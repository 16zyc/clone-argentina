from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = "/Users/a1111/Desktop/珠海/CLONE_AI运营指挥中心_功能说明书.docx"
NAVY = "0B2545"; BLUE = "1F4D78"; MINT = "18A984"; LIGHT = "E8EEF5"; PALE = "F4F7FA"; GOLD = "B7791F"; RED = "9B1C1C"; GRAY = "5C6B78"; WHITE = "FFFFFF"; BLACK = "17202A"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def set_col_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); grid.append(gc)
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn('w:tblW')); tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    for row in table.rows:
        for idx,cell in enumerate(row.cells):
            tcW=cell._tc.get_or_add_tcPr().find(qn('w:tcW')); tcW.set(qn('w:w'),str(widths[idx])); tcW.set(qn('w:type'),'dxa')

def add_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'; set_col_widths(t,widths)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_shading(c,NAVY); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE); r.font.size=Pt(9)
    set_repeat_table_header(t.rows[0])
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri%2: set_cell_shading(c,PALE)
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.12; r=p.add_run(str(val)); r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(BLACK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def add_bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.add_run(text); return p

def add_step(doc, num, title, body):
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_col_widths(t,[720,8640]); t.autofit=False
    c1,c2=t.rows[0].cells; set_cell_shading(c1,MINT); set_cell_margins(c1,100,80,100,80); set_cell_margins(c2,80,160,80,100)
    p=c1.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(num)); r.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE); r.font.size=Pt(14)
    p=c2.paragraphs[0]; r=p.add_run(title); r.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run('\n'+body).font.size=Pt(9)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def callout(doc, label, text, color=MINT):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_col_widths(t,[9360]); c=t.cell(0,0); set_cell_shading(c,PALE); set_cell_margins(c,140,180,140,180)
    p=c.paragraphs[0]; r=p.add_run(label+'  '); r.bold=True; r.font.color.rgb=RGBColor.from_string(color); r.font.size=Pt(9)
    r=p.add_run(text); r.font.size=Pt(9.5); r.font.color.rgb=RGBColor.from_string(BLACK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'等线'); normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(BLACK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Heading 1',16,NAVY,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,BLUE,10,5)]:
    s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'等线'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
styles['List Bullet'].font.name='Calibri'; styles['List Bullet']._element.rPr.rFonts.set(qn('w:eastAsia'),'等线'); styles['List Bullet'].font.size=Pt(10.5); styles['List Bullet'].paragraph_format.left_indent=Inches(.375); styles['List Bullet'].paragraph_format.first_line_indent=Inches(-.188); styles['List Bullet'].paragraph_format.space_after=Pt(4); styles['List Bullet'].paragraph_format.line_spacing=1.25

header=sec.header.paragraphs[0]; header.text='CLONE · Centro de Operaciones IA'; header.style=styles['Normal']; header.runs[0].font.size=Pt(8); header.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('功能说明书 · 2026.08').font.size=Pt(8)

# Cover - editorial_cover pattern
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('PRODUCT GUIDE'); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(MINT); p.paragraph_format.space_after=Pt(16)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CLONE'); r.bold=True; r.font.size=Pt(34); r.font.color.rgb=RGBColor.from_string(NAVY); p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('AI 运营指挥中心'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string(BLUE); p.paragraph_format.space_after=Pt(8)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Centro de Operaciones IA'); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(GRAY); p.paragraph_format.space_after=Pt(22)
callout(doc,'产品主张','不是更多看板，而是一支懂你业务、会协作、能执行且可追溯的 AI 团队。')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30); r=p.add_run('DEONCE B2B 运营层演示原型\n版本 1.0 · 2026 年 8 月'); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(GRAY)
doc.add_page_break()

doc.add_heading('1. 文档概述',level=1)
doc.add_paragraph('本说明书用于介绍 CLONE 西班牙语高保真原型的产品定位、功能结构、交互流程、演示方法以及与真实系统对接时的数据和安全边界。它适合比赛申报、项目路演、学生分工与后续 MVP 迭代。')
add_table(doc,['项目','说明'],[['产品名称','CLONE - Commercial Logic & Operation Native Engine'],['面向对象','阿根廷及南美 B2B 批发商户'],['原型语言','西班牙语（es-AR）'],['主要场景','AI 自动运营的一天：监控、分析、决策、审批、执行、审计'],['当前阶段','可点击高保真前端演示，不连接真实业务 API']], [1800,7560])
doc.add_heading('1.1 与 DEONCE 的关系',level=2)
doc.add_paragraph('CLONE 不取代 DEONCE，而是嵌入其上方的 AI 运营层。DEONCE 继续承担商品、客户、订单和交易；CLONE 负责读取经营信号、协调多个 Agent、生成可审批决策，并将执行结果与审计记录回写平台。')
callout(doc,'核心闭环','DEONCE 业务对象 → CLONE 多 Agent 分析 → 人工审批 → 执行与回写 → 可追溯记录')

doc.add_heading('2. 产品价值与技术差异',level=1)
add_table(doc,['对比维度','传统跨境系统','CLONE'],[['交互方式','菜单、报表、筛选条件','自然语言 + 可视化决策卡'],['工作单位','单个功能模块','多 Agent 共同完成业务目标'],['产出','数据和图表','建议、审批、执行与审计闭环'],['经验利用','人工操作流程','将商户偏好与白城批发经验键入 Agent'],['风险控制','权限与日志','证据链、人工审批、前后值和执行结果']], [1700,3300,4360])
add_bullet(doc,'经营者看到的不是“AI 说了什么”，而是“AI 发现了什么、为什么这样判断、建议做什么、最后是否执行成功”。')
add_bullet(doc,'低风险任务可自动化，价格、采购、税务和客户承诺等高风险操作默认要求人工审批。')
add_bullet(doc,'数据、Agent 和决策记录可持久化，让不同班次和不同团队成员共享业务记忆。')

doc.add_heading('3. 信息架构',level=1)
add_table(doc,['主视图','主要任务','现场演示价值'],[['Centro de mando','查看今日 KPI、时间线、异常与首要决策','一屏理解 AI 夜间已做什么'],['Operaciones','查看 10 大运营环节的任务与信号','证明能力不是单点工具'],['Equipo IA','查看调度关系、Agent 活动与自定义入口','展示业务经验克隆与多 Agent 协作'],['Aprobaciones','审核采购、促销和客户回复','展示人在回路和风险边界'],['Informes','查看日报、周报、导出预览和审计事件','展示可追溯、可交付和可复盘']], [1900,3730,3730])

doc.add_heading('4. 核心功能说明',level=1)
doc.add_heading('4.1 运营指挥中心',level=2)
add_bullet(doc,'四项首要 KPI：今日销售、AI 完成任务、节省工时、预计经营影响。')
add_bullet(doc,'按 Noche / Mañana / Tarde 分组的全天任务时间线，一次看到已完成、待处理和已计划任务。')
add_bullet(doc,'三个事件可下钻：TERMO-750 库存风险、竞品降价、ORG-4N 包装差评。')
doc.add_heading('4.2 补货决策闭环',level=2)
for n,t,b in [(1,'Detectar 检测','库存 Agent 发现可用库存只能覆盖 8.6 天。'),(2,'Analizar 分析','供应链 Agent 比较 1688 供应商，合规 Agent 校验 HS 9617.00 和 AFIP 风险。'),(3,'Recomendar 建议','调度器建议当日下单 1,800 件，投资 USD 14,760。'),(4,'Aprobar 审批','经营者可批准、调整数量或拒绝。'),(5,'Ejecutar 执行','原型模拟已批准、已计划、执行中、已执行的过渡。'),(6,'Auditar 审计','操作人、时间、Agent、证据、前后值和结果进入审计表。')]: add_step(doc,n,t,b)

doc.add_heading('4.3 十大运营环节',level=2)
ops=[('1','订单与绩效','订单、销售、毛利和异常对账'),('2','Listing 状态','上架状态、价格、类目和变体完整性'),('3','竞品监控','价格、排名、促销、评分和流量变化'),('4','评论与 Feedback','差评扫描、情绪分类、回复草稿与根因'),('5','库存与物流','库存水位、补货量、多仓调度与在途货件'),('6','多语内容','基于评论、搜索词和品牌调性生成内容'),('7','老品优化','销量趋势、词覆盖、评论变化和优化建议'),('8','退货分析','按 SKU、变体、批次和主题定位高频问题'),('9','新品开发','市场机会、供应商、利润、关税和 AI 评审'),('10','站外流量','站外渠道监控、归因和内容分发')]
add_table(doc,['#','环节','典型产出'],ops,[520,2400,6440])

doc.add_heading('4.4 专属 AI 团队',level=2)
doc.add_paragraph('团队页不只展示几个“机器人”，而是显示调度器与库存、供应链、合规、商业等 Agent 的职责边界。这种架构使任务可以并行执行，也使错误更易定位。')
add_table(doc,['键入维度','示例','系统结果'],[['业务目标','“优化批发老客户复购”','创建目标明确的商业 Agent'],['授权数据','DEONCE CRM + 历史订单','只在授权范围内读取数据'],['审批边界','任何客户优惠在发布前审批','生成人机分工规则'],['品牌/客户规则','老客户使用更简洁的西语语气','保持每个商户独特的经营方式']], [1900,3400,4060])

doc.add_heading('5. 固定演示数据',level=1)
add_table(doc,['字段','演示值','一致性说明'],[['SKU','TERMO-750','补货主案例'],['库存 / 日均销量','412 / 47.8','412 ÷ 47.8 = 8.6 天覆盖'],['总交期','24 天','远大于当前覆盖天数'],['建议量 / 单价','1,800 件 / USD 8.20','投资额 = USD 14,760'],['预计毛利率','36.2%','包含物流与税费的原型设定'],['保护销售额','ARS 18.4M','避免预计 16 天断货'],['竞品变化','ARS 42,900 → 39,990','降幅约 6.8%'],['差评对象','ORG-4N / L-0826','2 星，包装损坏，7 条同类问题']], [1850,2400,5110])
callout(doc,'数据声明','上述数字仅用于比赛演示和交互说明，不代表真实企业交易、库存、客户或税务记录。',GOLD)

doc.add_heading('6. 四分钟路演剧本',level=1)
script=[('0:00-0:40','首页','讲解 128 项夜间任务和四个 KPI','AI 已从“回答问题”进入“自主工作”'),('0:40-1:50','点击 TERMO-750','解释四 Agent 证据链，点击批准','数据、推理、人工决策和执行闭环'),('1:50-2:25','Equipo IA','展示协作图和 Forjar nuevo agente','可将每个商户的经营习惯键入 Agent'),('2:25-3:10','竞品降价','展示 AI 不建议盲目跟价','不只报警，而是保护利润的决策'),('3:10-3:40','Informes','查看审计记录和导出预览','每个动作可解释、可追溯'),('3:40-4:00','返回首页','用产品主张收尾','不是更多看板，而是懂业务的 AI 团队')]
add_table(doc,['时间','操作','讲解要点','评委记忆点'],script,[1150,1800,3200,3210])
add_bullet(doc,'开始前点击顶栏“Reiniciar demo”，确保审批、任务和审计记录回到初始状态。')
add_bullet(doc,'不要一开始讲模型参数；先讲经营问题、预计损失和人的决策权。')
add_bullet(doc,'如果时间被压缩，保留首页、补货审批和审计页三步即可完成闭环。')

doc.add_heading('7. 安全、合规与人机边界',level=1)
add_table(doc,['风险类型','原型处理','真实系统要求'],[['采购下单','仅模拟状态过渡','签名请求、金额阈值、双人审批和可撤销窗口'],['价格与促销','生成建议和审批项','渠道权限隔离、有效期和最低毛利保护'],['客户消息','草稿与模拟发送','隐私授权、敏感信息脱敏和语气审查'],['税务与海关','展示 HS 编码和风险结论','规则版本、来源签名、专业人员复核和生效日期'],['审计','展示操作人和状态','事件 ID、证据、前后值、模型/规则版本与不可篡改存储']], [1800,3100,4460])
callout(doc,'原型边界','当前页面明示“Entorno de demostración”。所有订单、客户消息、导出、税务和下单动作均为本地模拟，不会对外发送。',RED)

doc.add_heading('8. 本地运行与交互清单',level=1)
doc.add_paragraph('原型使用原生 HTML/CSS/JavaScript，无需安装前端依赖。为避免浏览器对本地文件的限制，建议在原型目录启动一个本地静态服务。')
callout(doc,'运行命令','python3 -m http.server 4173 --directory clone-prototype\n然后访问 http://localhost:4173/')
add_table(doc,['交互','位置','预期结果'],[['主导航','左侧边栏','切换五个主视图'],['事件下钻','首页时间线 / Operaciones','右侧打开证据与建议面板'],['补货审批','TERMO-750 详情','显示四段执行状态并更新审计'],['AI 对话','右下角 Preguntar a CLONE','点击预设问题或输入自定义问题'],['演示重置','顶栏 Reiniciar demo','所有临时状态返回初始值'],['导出预览','Informes','显示模拟文件名、格式与字段说明']], [1900,3200,4260])

doc.add_heading('9. 建议的学生分工',level=1)
add_table(doc,['角色','主要责任','阶段产出'],[['产品/路演','需求边界、演示剧本、评委问答','产品故事与 4 分钟讲稿'],['西语本地化','es-AR 术语、日期、金额与业务语气','西语 UI 文案审校表'],['前端','视图、响应式、交互状态与无障碍','可点击原型'],['数据/AI','演示数据字典、Agent 接口和推理证据','数据契约与模拟 API'],['安全/合规','审批阈值、审计字段、AFIP 规则来源','风险清单与演示边界']], [1900,3900,3560])

doc.add_heading('10. MVP 后续路线',level=1)
for n,t,b in [(1,'数据打通','优先接入 DEONCE 订单、SKU、库存和客户对象，建立统一 ID 与事件日志。'),(2,'两个真实闭环','先实现库存预警/补货草案以及差评分类/回复草稿，保留人工审批。'),(3,'多 Agent 调度','将数据获取、业务分析、合规检查与评审分成独立单元。'),(4,'审计与安全','完成规则版本、数据血缘、审批权限、失败回滚和敏感数据控制。'),(5,'扩展南美场景','在阿根廷验证后扩展葡萄牙语、巴西税务和多币种。')]: add_step(doc,n,t,b)

doc.add_paragraph()
callout(doc,'结语','CLONE 的先进性不应只用“大模型”表达，而应通过经营信号、多 Agent 协作、人工审批、真实执行和可追溯审计组成的完整业务闭环来证明。')

def iter_paragraphs(parent):
    for p in parent.paragraphs:
        yield p
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)

def set_east_asia_font(run):
    rpr=run._element.get_or_add_rPr(); fonts=rpr.rFonts
    if fonts is None:
        fonts=OxmlElement('w:rFonts'); rpr.insert(0,fonts)
    fonts.set(qn('w:ascii'),'SimSong')
    fonts.set(qn('w:hAnsi'),'SimSong')
    fonts.set(qn('w:eastAsia'),'宋体')
    fonts.set(qn('w:cs'),'SimSong')
    run.font.name='SimSong'

for p in iter_paragraphs(doc):
    for r in p.runs:
        set_east_asia_font(r)
for section in doc.sections:
    for part in (section.header, section.footer):
        for p in iter_paragraphs(part):
            for r in p.runs:
                set_east_asia_font(r)

doc.core_properties.title='CLONE AI 运营指挥中心功能说明书'; doc.core_properties.subject='西班牙语跨境 B2B AI 运营原型'; doc.core_properties.author='CLONE 项目组'
doc.save(OUT)
print(OUT)
